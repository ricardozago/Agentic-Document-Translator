import os
from typing import List
import docx
import pypdf

def read_document(file_path: str) -> str:
    """Reads content from .txt, .md, .pdf, or .docx files."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
        
    _, ext = os.path.splitext(file_path.lower())
    
    if ext in ['.txt', '.md']:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
            
    elif ext == '.pdf':
        text = []
        reader = pypdf.PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return "\n\n".join(text)
        
    elif ext == '.docx':
        doc = docx.Document(file_path)
        text = [paragraph.text for paragraph in doc.paragraphs]
        return "\n".join(text)
        
    else:
        raise ValueError(f"Formato de arquivo não suportado: {ext}")

def write_document(file_path: str, content: str) -> None:
    """Writes the translated content to a file, respecting format if possible."""
    _, ext = os.path.splitext(file_path.lower())
    
    if ext in ['.txt', '.md']:
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
            
    elif ext == '.docx':
        doc = docx.Document()
        paragraphs = content.split('\n')
        for p in paragraphs:
            if p.strip():
                doc.add_paragraph(p)
        doc.save(file_path)
        
    else:
        # Fallback to plain text if not supported
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

def split_text(text: str, max_chunk_size: int = 2000) -> List[str]:
    """
    Splits text into chunks of maximum size, trying to preserve paragraph and sentence boundaries.
    """
    if not text:
        return []
        
    # Heuristic to find paragraph delimiter
    if '\n\n' in text:
        paragraphs = text.split('\n\n')
        delimiter = '\n\n'
    else:
        paragraphs = text.split('\n')
        delimiter = '\n'
        
    chunks = []
    current_chunk = []
    current_len = 0
    
    for p in paragraphs:
        p_clean = p.strip()
        if not p_clean:
            continue
            
        p_len = len(p_clean)
        
        if p_len > max_chunk_size:
            # Paragraph is too large, split it by sentences
            if current_chunk:
                chunks.append(delimiter.join(current_chunk))
                current_chunk = []
                current_len = 0
            
            # Sentence splitter
            sentences = []
            current_sentence = []
            current_sent_len = 0
            words = p_clean.split(' ')
            for word in words:
                current_sentence.append(word)
                current_sent_len += len(word) + 1
                if word.endswith(('.', '!', '?')) and current_sent_len > 100:
                    sentences.append(" ".join(current_sentence))
                    current_sentence = []
                    current_sent_len = 0
            if current_sentence:
                sentences.append(" ".join(current_sentence))
                
            # Now chunk the sentences
            sub_chunk = []
            sub_len = 0
            for s in sentences:
                s_len = len(s)
                if sub_len + s_len + 1 > max_chunk_size:
                    if sub_chunk:
                        chunks.append(" ".join(sub_chunk))
                    sub_chunk = [s]
                    sub_len = s_len
                else:
                    sub_chunk.append(s)
                    sub_len += s_len + 1
            if sub_chunk:
                chunks.append(" ".join(sub_chunk))
                
        else:
            # Standard chunking
            if current_len + p_len + len(delimiter) > max_chunk_size:
                if current_chunk:
                    chunks.append(delimiter.join(current_chunk))
                current_chunk = [p_clean]
                current_len = p_len
            else:
                current_chunk.append(p_clean)
                current_len += p_len + len(delimiter)
                
    if current_chunk:
        chunks.append(delimiter.join(current_chunk))
        
    return chunks


def get_cache_path(file_path: str) -> str:
    """Retorna o caminho do arquivo de cache para o arquivo fornecido."""
    import os
    directory, filename = os.path.split(file_path)
    cache_filename = f".{filename}.translation_cache.json"
    return os.path.join(directory, cache_filename)


def get_file_hash(file_path: str) -> str:
    """Calcula o hash SHA256 do arquivo para detectar modificações."""
    import hashlib
    hasher = hashlib.sha256()
    with open(file_path, 'rb') as f:
        buf = f.read(65536)
        while len(buf) > 0:
            hasher.update(buf)
            buf = f.read(65536)
    return hasher.hexdigest()


def save_translation_cache(
    file_path: str,
    target_lang: str,
    source_lang: str,
    current_chunk_idx: int,
    translated_chunks: List[str],
    final_chunks: List[str],
    glossary: dict,
    chunks: List[str]
) -> None:
    """Salva o progresso atual da tradução em um arquivo JSON de cache."""
    import json
    cache_path = get_cache_path(file_path)
    try:
        file_hash = get_file_hash(file_path)
    except Exception:
        file_hash = ""
        
    cache_data = {
        "file_hash": file_hash,
        "target_lang": target_lang,
        "source_lang": source_lang,
        "current_chunk_idx": current_chunk_idx,
        "translated_chunks": translated_chunks,
        "final_chunks": final_chunks,
        "glossary": glossary,
        "chunks": chunks
    }
    
    with open(cache_path, 'w', encoding='utf-8') as f:
        json.dump(cache_data, f, ensure_ascii=False, indent=2)


def load_translation_cache(file_path: str, target_lang: str, source_lang: str) -> dict:
    """Carrega o progresso da tradução se o arquivo e as configurações coincidirem."""
    import os
    import json
    cache_path = get_cache_path(file_path)
    if not os.path.exists(cache_path):
        return None
        
    try:
        with open(cache_path, 'r', encoding='utf-8') as f:
            cache_data = json.load(f)
            
        # Validar integridade: se o arquivo de origem mudou, invalida o cache
        file_hash = get_file_hash(file_path)
        if cache_data.get("file_hash") != file_hash:
            return None
            
        # Validar se as opções de idioma coincidem
        if cache_data.get("target_lang") != target_lang or cache_data.get("source_lang") != source_lang:
            return None
            
        return cache_data
    except Exception:
        return None


def clear_translation_cache(file_path: str) -> None:
    """Remove o arquivo de cache ao concluir a tradução."""
    import os
    cache_path = get_cache_path(file_path)
    if os.path.exists(cache_path):
        try:
            os.remove(cache_path)
        except Exception:
            pass


def parse_vocabulary_content(content_bytes: bytes, file_name: str) -> dict:
    """Parses vocabulary from content bytes (JSON or TXT formats)."""
    import os
    import json
    _, ext = os.path.splitext(file_name.lower())
    vocab = {}
    if ext == '.json':
        try:
            data = json.loads(content_bytes.decode('utf-8'))
            if isinstance(data, dict):
                return {str(k).strip(): str(v).strip() for k, v in data.items()}
        except Exception as e:
            raise ValueError(f"Erro ao analisar JSON de vocabulário: {e}")
    else:
        try:
            lines = content_bytes.decode('utf-8').splitlines()
            for line in lines:
                line_clean = line.strip()
                if not line_clean or line_clean.startswith('#'):
                    continue
                if '=' in line_clean:
                    parts = line_clean.split('=', 1)
                elif ':' in line_clean:
                    parts = line_clean.split(':', 1)
                else:
                    continue
                k, v = parts[0].strip(), parts[1].strip()
                if k and v:
                    vocab[k] = v
        except Exception as e:
            raise ValueError(f"Erro ao analisar arquivo de vocabulário TXT: {e}")
    return vocab


def load_vocabulary(file_path: str) -> dict:
    """Loads vocabulary dictionary from a local file path."""
    import os
    if not file_path or not os.path.exists(file_path):
        return {}
    try:
        with open(file_path, 'rb') as f:
            content_bytes = f.read()
        return parse_vocabulary_content(content_bytes, os.path.basename(file_path))
    except Exception as e:
        raise ValueError(f"Não foi possível carregar o vocabulário de {file_path}: {e}")

